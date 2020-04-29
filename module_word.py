# Bibliotecas
import docx
import pandas as pd
import numpy as np
from tqdm import tqdm
import os
import Start
# a.tables[0].cell(l[0][0],l[0][1]).text  (esse é o segredo para encontrar a informação dentro da tabela docx)

 
# Funções

# Menu inicial
def module_init():
    os.chdir("C:/Users/ADM/Desktop/Temporários/Rats") # rats significa a pasta aonde estão os arquivos a serem analizados
    print("Seu diretório atual é: C:/Users/ADM/Desktop/Temporários/Rats" )
    r =input("\nDeseja mudar de diretório? [y/n]")
    try:
        if r =='y':
            return select_directory()
        elif r =='n':
            global directory
            directory = str(os.getcwd())
            return list_directorys()
        else:
            print("Comando inválido")
            return module_init()
    except KeyboardInterrupt:
        r =input("Deseja voltar ao menu inicial? [y/n]")
        if r == 'y':
            import Start
            Start.initialization() 
        else:
            return module_init()

def select_directory():
    global directory
    directory = input("Informe o diretório em que o arquivo se encontra:")
    os.chdir(directory)
    print("Done")
    list_directorys()


def list_directorys():
    arquivos = os.listdir(directory)
    global df_files
    df_files = pd.DataFrame(arquivos, columns=['Arquivos'])
    print(df_files)
    return File()

def File():
    try:
        File = df_files.iloc[int(input("\nInsira o índice do arquivo:"))].to_string()
        global doc
        doc = File.rsplit("Arquivos    ")[1]
        if doc[-5:] == '.docx':
            global a
            a =docx.Document(doc)
            print("\nForam encontradas "+str(len(a.tables))+ " tabelas no documento\n")
            choose_function()
        else:
            print("\nO arquivo deve ser de extensão .docx\n")
            return list_directorys()
    except NameError:
        print("Documento não encotrado\n")
        return File()

# Escolha das funções e determinação de intervalo de scrapping
def choose_function():
    global i, lim , dif, scr, error
    error =0
    print(df_scrapping)
    scr = input("\nSelecione uma opção: ")
    try:
        if scr == '0':
            print("Opção 0 escolhida")
            i = 0
            lim = len(a.tables)
            dif = lim - i
            r=""
            while r!= 'y'or r !='n':
                r = input("\nComeçar scrapping? [y/n]")
                if r == 'y':
                    scrapping_tables()
                    break
                elif r == 'n':
                    print("Ok doutor")
                    return choose_function()            
                else:
                    print("Comando inválido")
        elif scr == '1':
            print("Opção 1 escolhida")
            i = int(input("Selecione o limite inferior: "))
            if i > len(a.tables):
                raise Exception("Limite definido é maior que o total de tabelas do documento")
            if i <0:
                raise Exception("Limmite não pode ser inferior a zero")
            lim = int(input("Selecione o limite superior: "))
            if lim < i:
                raise Exception("O limite superior deve ser maior que o limite inferior")
            if lim > len(a.tables):
                raise Exception("Limite definido é maior que o total de tabelas do documento")   
            dif = lim - i
            r=""
            while r!= 'y'or r !='n':
                r = input("\nComeçar scrapping? [y/n]")
                if r == 'y':
                    scrapping_tables()
                    break
                elif r == 'n':
                    print("Ok doutor")
                    return choose_function()            
                else:
                    print("Comando inválido")
        elif scr == '2':
            print("Opção 2 escolhida")
            i = int(input("Informe o número da tabela: "))
            r=""
            while r!= 'y'or r !='n':
                r = input("\nComeçar scrapping? [y/n]")
                if r == 'y':
                    scrapping_table()
                    break                                           
                elif r == 'n':
                    print("Ok doutor")
                    return choose_function() 
        else:
            print("\nComando inválido, selecione uma opção do índice da tabela\n")
            return choose_function()

    except KeyboardInterrupt:
        r =''
        while r != 'f' or r != 'i':
            r = input("\nDeseja voltar às funções [f] ou a seleção de documento [i]? [f/i]")
            if r == 'f':
                choose_function()
            elif r =='i':
                File()
            else:
                print("\nComando inválido")

# Scrapping de 1 tabela
def scrapping_table():
    try:
        j=0
        a.tables[i]
        array = np.empty((len(a.tables[i].rows),len(a.tables[i].columns)))    
        l=[]
        it = np.nditer(array, flags=['multi_index'])
        while not it.finished:
            l.append(it.multi_index)
            it.iternext()
            array = array.astype('object')
        while j < len(l):
            array[l[j][0],l[j][1]] = a.tables[i].cell(l[j][0],l[j][1]).text       
            j+=1
        df = pd.DataFrame(data=array[1:],columns=array[0,:])
        df.to_excel('C:/Users/ADM/Desktop/Temporários/Output_Word_module/tabelas.xlsx',sheet_name='Tabela selecionada')
        results()
                                                    
    except PermissionError:
        input("\nPermission Error:\nVerifique se o arquivo EXCEL está aberto, favor fechá-lo para prosseguir")
        return choose_function()

# Scrapping de várias tabelas
def scrapping_tables():
        global i, error
        try:
            while i < lim:
                for i in tqdm(range(i,lim)):
                    j=0
                    a.tables[i] 
                    array = np.empty((len(a.tables[i].rows),len(a.tables[i].columns)))    
                    l=[]
                    it = np.nditer(array, flags=['multi_index'])
                    while not it.finished:
                        l.append(it.multi_index)
                        it.iternext()
                        array = array.astype('object')
                    while j < len(l):
                        array[l[j][0],l[j][1]] = a.tables[i].cell(l[j][0],l[j][1]).text       
                        j+=1
                    df = pd.DataFrame(data=array[1:],columns=array[0,:])
                    if i == 0:
                        df.to_excel('C:/Users/ADM/Desktop/Temporários/Output_Word_module/tabelas.xlsx',sheet_name='Tabela 0')
                        i+=1
                    else:
                        with pd.ExcelWriter('C:/Users/ADM/Desktop/Temporários/Output_Word_module/tabelas.xlsx',mode='a') as writer:
                            df.to_excel(writer,sheet_name='Tabela '+str(i))
                        i+=1
                results()
        except PermissionError:
            input("\nPermission Error:\nO arquivo EXCEL está aberto, favor fechá-lo para prosseguir")        
            return scrapping_tables()
        except IndexError:
            i+=1
            error+=1
            print("Tabela "+ str(i) +" falhou" )
            return scrapping_tables()

# Resultado de Scrappings
def results():
    if scr != '2':
        print("\n"+str(dif-error)+ " tabelas foram convertidas com sucesso")
    else:
        print("\nTabela carregada com sucesso")
    r=""
    while r != 'y' or r != 'n':
        r = input("\nDeseja fazer scrapping em outro documento? [y/n]")
        if r == 'y':
            File()
        elif r == 'n':
            print("Até mais campeão!")
            Start.initialization()
        else:
            print("Comando inválido")

# Execução

df_scrapping = pd.DataFrame(data=['Todas as tabelas','Intervalo de tabelas','Tabela específica'],columns=['Tipos de scrapping'])
module_init()


